from docx import Document
from docx.shared import Inches
import re
import os

def markdown_to_docx(markdown_file, docx_file):
    document = Document()

    with open(markdown_file, 'r', encoding='utf-8') as f:
        markdown_content = f.read()

    lines = markdown_content.split('\n')

    # Keep track of the last added table to append rows
    document._last_table = None

    for line in lines:
        line = line.strip()
        if not line:
            document.add_paragraph()
            continue

        # Headings
        if line.startswith('### '):
            document.add_heading(line[4:], level=3)
            document._last_table = None # Reset table tracking after a heading
        elif line.startswith('## '):
            document.add_heading(line[3:], level=2)
            document._last_table = None
        elif line.startswith('# '):
            document.add_heading(line[2:], level=1)
            document._last_table = None
        # Lists
        elif line.startswith('* '):
            document.add_paragraph(line[2:], style='List Bullet')
            document._last_table = None
        # Tables
        elif line.startswith('|'):
            # Basic table parsing - assumes header and separator
            if '| :---' in line or '| ---' in line: # Separator line
                continue
            
            cells = [c.strip() for c in line.split('|') if c.strip()]
            
            # If no table is being tracked or the number of columns changed, create a new table
            if not document._last_table or len(document._last_table.columns) != len(cells):
                document._last_table = document.add_table(rows=1, cols=len(cells))
                document._last_table.style = 'Table Grid'
                # Add header row
                hdr_cells = document._last_table.rows[0].cells
                for i, cell_text in enumerate(cells):
                    hdr_cells[i].text = cell_text
            else:
                # Add row to existing table
                row_cells = document._last_table.add_row().cells
                for i, cell_text in enumerate(cells):
                    row_cells[i].text = cell_text
        # Screenshot placeholders
        elif line.startswith('[Screenshot:'):
            p = document.add_paragraph()
            p.add_run(line).bold = True
            document.add_paragraph("    [Insert your screenshot here, e.g., 6 inches wide]")
            document.add_paragraph() # Add some space after placeholder
            document._last_table = None
        # Bold text (simple)
        elif re.match(r'\*\*(.*?)\*\*', line):
            p = document.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
            document._last_table = None
        # Paragraphs
        else:
            document.add_paragraph(line)
            document._last_table = None
    
    document.save(docx_file)
    print(f"Report saved to {docx_file}")

if __name__ == "__main__":
    markdown_to_docx('report.md', 'Yarra_Project_Report.docx')
